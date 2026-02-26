from fastapi import FastAPI, HTTPException
from docx import Document
import httpx
import os
from datetime import datetime

app = FastAPI()

# === Bitrix24 настройки ===
WEBHOOK = "https://izyskaniya.bitrix24.ru/rest/13614/rj3pqolk1fiu6hfr/"
DISK_FOLDER_ID = "1706930"

def format_cost(cost_str):
    try:
        return f"{int(cost_str):,}".replace(",", " ")
    except (ValueError, TypeError):
        return "0"

def remove_section(doc, start_text, end_text):
    """Удаляет все параграфы между start_text и end_text (включительно)"""
    paragraphs_to_remove = []
    in_section = False
    for para in doc.paragraphs:
        if start_text in para.text:
            in_section = True
        if in_section:
            paragraphs_to_remove.append(para)
        if end_text in para.text and in_section:
            in_section = False
            break
    for para in paragraphs_to_remove:
        p = para._element
        p.getparent().remove(p)

@app.get("/generate-kp")
async def generate_kp(
    object_name: str,
    address: str,
    cadastral_number: str = "",
    date: str = None,
    total_cost: str = "0",
    advance_percent: str = "50",
    validity_days: str = "30",
    # ИГИ
    igi: str = "0",
    igi_drilling_depth: str = "5",
    igi_boreholes: str = "4",
    igi_sounding_points: str = "4",
    igi_duration_days: str = "35",
    igi_cost: str = "0",
    # ИГДИ
    igdi: str = "0",
    igdi_area_ha: str = "0",
    igdi_scale: str = "1:500",
    igdi_contour_interval: str = "0.5",
    igdi_duration_days: str = "45",
    igdi_survey_days: str = "10",
    igdi_coordination_days: str = "35",
    igdi_cost: str = "0",
    igdi_survey_cost: str = "0",
    igdi_coordination_cost: str = "0",
    igdi_report_cost: str = "0",
    # ИЭИ
    iei: str = "0",
    iei_area_ha: str = "0",
    iei_gamma_points: str = "0",
    iei_noise_points: str = "0",
    iei_emi_points: str = "0",
    iei_soil_samples: str = "0",
    iei_bio_samples: str = "0",
    iei_rad_samples: str = "0",
    iei_surface_water_samples: str = "0",
    iei_sediment_samples: str = "0",
    iei_water_samples: str = "0",
    iei_water_boreholes: str = "0",
    iei_layered_samples_deep: str = "0",
    iei_deep_boreholes: str = "0",
    iei_layered_samples_shallow: str = "0",
    iei_shallow_boreholes: str = "0",
    iei_background_soil_samples: str = "0",
    iei_agro_samples: str = "0",
    iei_pits: str = "0",
    iei_duration_days: str = "35",
    iei_cost: str = "0",
    # ИГМИ
    igmi: str = "0",
    igmi_route_km: str = "0",
    igmi_photo_count: str = "0",
    igmi_wind_rose_count: str = "0",
    igmi_duration_days: str = "40",
    igmi_cost: str = "0"
):
    try:
        # Подготовка данных
        data = {
            "object_name": object_name,
            "address": address,
            "cadastral_number": cadastral_number or "—",
            "date": date or datetime.now().strftime("%d.%m.%Y"),
            "total_cost": format_cost(total_cost),
            "advance_percent": advance_percent,
            "validity_days": validity_days,
            "igi": igi == "1",
            "igi_drilling_depth": igi_drilling_depth,
            "igi_boreholes": igi_boreholes,
            "igi_sounding_points": igi_sounding_points,
            "igi_duration_days": igi_duration_days,
            "igi_cost": format_cost(igi_cost),
            "igdi": igdi == "1",
            "igdi_area_ha": igdi_area_ha,
            "igdi_scale": igdi_scale,
            "igdi_contour_interval": igdi_contour_interval,
            "igdi_duration_days": igdi_duration_days,
            "igdi_survey_days": igdi_survey_days,
            "igdi_coordination_days": igdi_coordination_days,
            "igdi_cost": format_cost(igdi_cost),
            "igdi_survey_cost": format_cost(igdi_survey_cost),
            "igdi_coordination_cost": format_cost(igdi_coordination_cost),
            "igdi_report_cost": format_cost(igdi_report_cost),
            "iei": iei == "1",
            "iei_area_ha": iei_area_ha,
            "iei_gamma_points": iei_gamma_points,
            "iei_noise_points": iei_noise_points,
            "iei_emi_points": iei_emi_points,
            "iei_soil_samples": iei_soil_samples,
            "iei_bio_samples": iei_bio_samples,
            "iei_rad_samples": iei_rad_samples,
            "iei_surface_water_samples": iei_surface_water_samples,
            "iei_sediment_samples": iei_sediment_samples,
            "iei_water_samples": iei_water_samples,
            "iei_water_boreholes": iei_water_boreholes,
            "iei_layered_samples_deep": iei_layered_samples_deep,
            "iei_deep_boreholes": iei_deep_boreholes,
            "iei_layered_samples_shallow": iei_layered_samples_shallow,
            "iei_shallow_boreholes": iei_shallow_boreholes,
            "iei_background_soil_samples": iei_background_soil_samples,
            "iei_agro_samples": iei_agro_samples,
            "iei_pits": iei_pits,
            "iei_duration_days": iei_duration_days,
            "iei_cost": format_cost(iei_cost),
            "igmi": igmi == "1",
            "igmi_route_km": igmi_route_km,
            "igmi_photo_count": igmi_photo_count,
            "igmi_wind_rose_count": igmi_wind_rose_count,
            "igmi_duration_days": igmi_duration_days,
            "igmi_cost": format_cost(igmi_cost)
        }

        # Загрузка шаблона
        template_path = os.path.join(os.path.dirname(__file__), "templates", "kp_template.docx")
        if not os.path.exists(template_path):
            raise HTTPException(status_code=400, detail="Шаблон КП не найден")

        doc = Document(template_path)

        # Замена меток
        for paragraph in doc.paragraphs:
            text = paragraph.text
            for key, value in data.items():
                if isinstance(value, bool):
                    continue
                placeholder = f"{{{{{key}}}}}"
                if placeholder in text:
                    text = text.replace(placeholder, str(value))
            paragraph.text = text

        # Удаление ненужных разделов
        if not data["igi"]:
            remove_section(doc, "1. Комплекс инженерно-геологических изысканий:", "Стоимость работ по п.1 составит:")
        if not data["igdi"]:
            remove_section(doc, "2. Комплекс инженерно-геодезических изысканий:", "* В составе работ подача в ИСОГД")
        if not data["iei"]:
            remove_section(doc, "3. Комплекс инженерно-экологических изысканий:", "биотестирование почво-грунтов с целью их экотоксикологической оценки")
        if not data["igmi"]:
            remove_section(doc, "4. Комплекс гидрометеорологических изысканий:", "*в стоимость работ не включены:")

        # Сохранение файла
        safe_name = "".join(c for c in object_name if c.isalnum() or c in " _-")
        filename = f"KP_{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx"
        output_path = f"/tmp/{filename}"
        doc.save(output_path)

        # Загрузка в Bitrix24
        async with httpx.AsyncClient(timeout=30) as client:
            prep_resp = await client.post(
                f"{WEBHOOK}disk.folder.uploadfile.json",
                data={"id": DISK_FOLDER_ID}
            )
            prep_data = prep_resp.json()
            if "result" not in prep_data or "uploadUrl" not in prep_data["result"]:
                raise HTTPException(status_code=500, detail="Не удалось получить uploadUrl от Bitrix24")

            upload_url = prep_data["result"]["uploadUrl"]
            field_name = prep_data["result"].get("field", "file")

            with open(output_path, "rb") as f:
                files = {field_name: (filename, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
                upload_resp = await client.post(upload_url, files=files)

            upload_result = upload_resp.json()
            if "result" not in upload_result:
                raise HTTPException(status_code=500, detail="Ошибка загрузки файла в Bitrix24")

            file_id = str(upload_result["result"]["ID"])

        # Формирование ссылки
        download_url = f"https://izyskaniya.bitrix24.ru/disk/showFile/{file_id}/?filename={filename}"

        # Удаление временного файла
        os.remove(output_path)

        return {
            "status": "success",
            "message": f"📄 КП готов! Скачать можно по ссылке: {download_url}",
            "download_url": download_url
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))